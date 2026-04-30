from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import io
import json

load_dotenv()

app = FastAPI(title="AI Career Copilot API", version="1.0.0")

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, change to frontend url
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize OpenAI client
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    raise ValueError("OPENAI_API_KEY environment variable is required")
client = OpenAI(api_key=OPENAI_API_KEY)

class ResumeGenerationRequest(BaseModel):
    full_name: str
    job_title: str
    role: str
    company: str = ""
    years_experience: int = 0
    skills: list[str] = []
    achievements: list[str] = []

@app.get("/")
async def root():
    return {"message": "Welcome to AI Career Copilot API"}

@app.post("/generate-resume")
async def generate_resume(request: ResumeGenerationRequest):
    """Generate ATS-friendly resume using OpenAI"""
    try:
        prompt = f"""Create a professional, ATS-friendly resume for the following person:
        
Name: {request.full_name}
Target Job Title: {request.job_title}
Role: {request.role}
Company: {request.company if request.company else 'Multiple'}
Years of Experience: {request.years_experience}
Skills: {', '.join(request.skills) if request.skills else 'General technical skills'}
Key Achievements: {', '.join(request.achievements) if request.achievements else 'Strong problem-solver and team player'}

Please generate:
1. A professional summary (2-3 lines)
2. Core technical skills (10-15 skills)
3. Professional experience section with 2-3 roles including:
   - Job title and company
   - Duration (realistic based on years of experience)
   - 3-4 achievement bullets with metrics where possible
4. Education section
5. Certifications (if applicable)

Make it ATS-friendly by:
- Using clear section headers
- Including industry keywords
- Using quantifiable metrics
- Keeping formatting simple and clean
- Using standard fonts

Format the response as JSON with keys: summary, skills, experience, education, certifications"""

        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=2000
        )

        content = response.choices[0].message.content
        
        # Try to parse as JSON, fallback to structured data
        try:
            resume_data = json.loads(content)
        except:
            # If not valid JSON, create structured data
            resume_data = {
                "summary": "Experienced professional with strong technical skills and proven track record of success.",
                "skills": request.skills or ["Python", "JavaScript", "React", "Node.js", "SQL", "REST APIs"],
                "experience": [{
                    "title": request.job_title,
                    "company": request.company or "Tech Company",
                    "duration": f"{request.years_experience} years",
                    "bullets": ["Delivered high-impact projects", "Led cross-functional teams", "Improved efficiency and performance"]
                }],
                "education": ["Bachelor's Degree in Computer Science"],
                "certifications": []
            }

        return {
            "status": "success",
            "resume_data": resume_data,
            "name": request.full_name,
            "job_title": request.job_title
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate-resume-docx")
async def generate_resume_docx(request: ResumeGenerationRequest):
    """Generate and return DOCX resume file"""
    try:
        # First generate resume content
        prompt = f"""Create a professional, ATS-friendly resume for:
Name: {request.full_name}
Job Title: {request.job_title}
Role: {request.role}
Company: {request.company or 'Tech Company'}
Experience: {request.years_experience} years
Skills: {', '.join(request.skills) if request.skills else 'Technical skills'}

Format as JSON with: summary, skills (array), experience (array with title, company, duration, bullets), education (array), certifications (array)"""

        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=2000
        )

        content = response.choices[0].message.content
        
        try:
            resume_data = json.loads(content)
        except:
            resume_data = {
                "summary": "Experienced professional ready to contribute to your organization.",
                "skills": request.skills or ["Python", "JavaScript", "React"],
                "experience": [{"title": request.job_title, "company": request.company or "Company", "duration": "Ongoing", "bullets": ["Key achievement 1", "Key achievement 2"]}],
                "education": ["Bachelor's Degree"],
                "certifications": []
            }

        # Create DOCX document
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Header - Name and Contact
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(request.full_name)
        name_run.font.size = Pt(16)
        name_run.font.bold = True
        name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        title_para = doc.add_paragraph(f"{request.job_title} | {request.role}")
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_para_format = title_para.paragraph_format
        title_para_format.space_before = Pt(0)
        title_para_format.space_after = Pt(6)

        # Professional Summary
        summary_heading = doc.add_heading('PROFESSIONAL SUMMARY', level=2)
        summary_heading.paragraph_format.space_before = Pt(6)
        summary_heading.paragraph_format.space_after = Pt(3)
        doc.add_paragraph(resume_data.get('summary', 'Professional with strong technical expertise.'))

        # Skills
        skills_heading = doc.add_heading('CORE SKILLS', level=2)
        skills_heading.paragraph_format.space_before = Pt(6)
        skills_heading.paragraph_format.space_after = Pt(3)
        skills_text = ', '.join(resume_data.get('skills', []))
        doc.add_paragraph(skills_text)

        # Experience
        exp_heading = doc.add_heading('PROFESSIONAL EXPERIENCE', level=2)
        exp_heading.paragraph_format.space_before = Pt(6)
        exp_heading.paragraph_format.space_after = Pt(3)

        for job in resume_data.get('experience', []):
            job_para = doc.add_paragraph()
            job_title_run = job_para.add_run(job.get('title', 'Job Title'))
            job_title_run.bold = True
            job_para.add_run(f" | {job.get('company', 'Company')}")
            job_para.add_run(f" | {job.get('duration', '2023-2024')}")
            job_para.paragraph_format.space_after = Pt(3)

            for bullet in job.get('bullets', []):
                doc.add_paragraph(bullet, style='List Bullet')

        # Education
        if resume_data.get('education'):
            edu_heading = doc.add_heading('EDUCATION', level=2)
            edu_heading.paragraph_format.space_before = Pt(6)
            edu_heading.paragraph_format.space_after = Pt(3)
            for edu in resume_data.get('education', []):
                doc.add_paragraph(edu)

        # Certifications
        if resume_data.get('certifications'):
            cert_heading = doc.add_heading('CERTIFICATIONS', level=2)
            cert_heading.paragraph_format.space_before = Pt(6)
            cert_heading.paragraph_format.space_after = Pt(3)
            for cert in resume_data.get('certifications', []):
                doc.add_paragraph(cert)

        # Save to bytes
        bio_output = io.BytesIO()
        doc.save(bio_output)
        bio_output.seek(0)

        filename = f"{request.full_name.replace(' ', '_')}_Resume.docx"
        
        return FileResponse(
            io.BytesIO(bio_output.getvalue()),
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=filename
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
